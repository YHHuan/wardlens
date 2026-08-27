from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path

from wardlens.models import PatientBundle, PatientSummary
from wardlens.services.summary import extract_top_todos, recent_reports


def export_patient_list_csv(patients: list[PatientSummary], path: Path) -> Path:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Ward", "Bed", "Name", "HistNo", "Age", "Sex", "CaseNo"])
        for patient in patients:
            writer.writerow(
                [
                    patient.ward,
                    patient.bed,
                    patient.name,
                    patient.histno,
                    patient.age,
                    patient.sex,
                    patient.case_no,
                ]
            )
    return path


def export_patient_list_docx(
    patients: list[PatientSummary],
    bundles: dict[str, PatientBundle],
    path: Path,
    *,
    doctor_id: str = "",
) -> Path:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("DOCX export dependency is unavailable.") from exc

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)

    title = document.add_paragraph()
    title.add_run("WardLens 查房清單｜機密醫療資料").bold = True
    title.add_run(f"｜{datetime.now().astimezone().strftime('%Y-%m-%d %H:%M')}")
    if doctor_id:
        title.add_run(f"｜醫師燈號 {doctor_id}")
    notice = document.add_paragraph("內容為唯讀擷取／AI 草稿輔助；請逐項核對院內原始病歷。")
    notice.runs[0].italic = True

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["病人資料", "Admission TODO", "近期／Pending 報告", "擷取狀態"]
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for patient in patients:
        cells = table.add_row().cells
        cells[0].text = "\n".join(
            value
            for value in (
                patient.location,
                patient.name,
                f"MRN {patient.histno}",
                " ".join(part for part in (patient.age, patient.sex) if part),
            )
            if value
        )
        bundle = bundles.get(patient.local_key)
        if bundle is None:
            cells[1].text = "尚未載入；不可解讀為無待辦。"
            cells[2].text = "尚未載入；不可解讀為無新報告。"
            cells[3].text = "只有基本清單"
        else:
            cells[1].text = "\n".join(extract_top_todos(bundle)) or "未辨識，請核對原文。"
            cells[2].text = "\n".join(recent_reports(bundle)) or "未取得報告索引。"
            cells[3].text = (
                f"截止 {bundle.fetched_at.astimezone().strftime('%H:%M')}\n"
                f"來源 {len(bundle.records)}｜警告 {len(bundle.warnings)}"
            )

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.save(path)
    return path
